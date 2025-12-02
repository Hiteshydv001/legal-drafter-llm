import os
import uuid
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
from reportlab.lib import colors
from app.schemas.models import LegalDocument
from app.core.config import settings
from app.core.logger import get_logger
from app.core.exceptions import DocumentCreationError

logger = get_logger("Doc_Service")

class DocumentService:
    @staticmethod
    def create_docx(data: LegalDocument) -> str:
        try:
            doc = Document()
            
            # --- Global Styles ---
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(11)

            # 1. Title
            title_para = doc.add_heading(data.title.upper(), level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Add Bold and Underline to title run
            for run in title_para.runs:
                run.bold = True
                run.underline = True
            
            # 2. Intro
            intro = doc.add_paragraph(data.introductory_clause)
            intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            doc.add_paragraph() 

            # 3. Clauses
            for idx, clause in enumerate(data.clauses, 1):
                h = doc.add_heading(f"{idx}. {clause.heading}", level=2)
                p = doc.add_paragraph(clause.content)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(12)

            # 4. Signatures
            doc.add_page_break()
            doc.add_heading("IN WITNESS WHEREOF", level=3)
            doc.add_paragraph("The Parties have executed this Agreement on the date first written above.\n")

            table = doc.add_table(rows=1, cols=2)
            table.autofit = True
            
            row_cells = table.rows[0].cells
            for i, signer in enumerate(data.signature_blocks):
                # Add new row if current is full
                if i > 0 and i % 2 == 0:
                    row_cells = table.add_row().cells
                
                target_cell = row_cells[i % 2]
                p = target_cell.paragraphs[0]
                p.add_run(f"\n\n__________________________\n{signer}").bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Save File
            filename = f"Legal_Doc_{uuid.uuid4().hex[:8]}.docx"
            file_path = os.path.join(settings.OUTPUT_DIR, filename)
            doc.save(file_path)
            
            logger.info(f"Document saved successfully: {file_path}")
            return file_path

        except Exception as e:
            logger.error(f"Docx creation failed: {str(e)}")
            raise DocumentCreationError(f"File generation failed: {str(e)}")

    @staticmethod
    def create_pdf(data: LegalDocument) -> str:
        try:
            filename = f"Legal_Doc_{uuid.uuid4().hex[:8]}.pdf"
            file_path = os.path.join(settings.OUTPUT_DIR, filename)
            
            doc = SimpleDocTemplate(file_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            # Custom Styles
            title_style = ParagraphStyle(
                'Title',
                parent=styles['Heading1'],
                alignment=TA_CENTER,
                fontSize=16,
                spaceAfter=20
            )
            
            normal_style = ParagraphStyle(
                'Normal_Justified',
                parent=styles['Normal'],
                alignment=TA_JUSTIFY,
                fontSize=11,
                leading=14,
                spaceAfter=12
            )

            heading_style = ParagraphStyle(
                'Heading_Left',
                parent=styles['Heading2'],
                fontSize=13,
                spaceAfter=10,
                spaceBefore=10
            )

            # 1. Title
            story.append(Paragraph(data.title.upper(), title_style))
            story.append(Spacer(1, 12))

            # 2. Intro
            story.append(Paragraph(data.introductory_clause, normal_style))
            story.append(Spacer(1, 12))

            # 3. Clauses
            for idx, clause in enumerate(data.clauses, 1):
                story.append(Paragraph(f"{idx}. {clause.heading}", heading_style))
                story.append(Paragraph(clause.content, normal_style))

            # 4. Signatures
            story.append(Spacer(1, 30))
            story.append(Paragraph("IN WITNESS WHEREOF", heading_style))
            story.append(Paragraph("The Parties have executed this Agreement on the date first written above.", normal_style))
            story.append(Spacer(1, 20))

            # Signature Table
            sig_data = []
            row = []
            for i, signer in enumerate(data.signature_blocks):
                cell_text = f"\n\n__________________________\n{signer}"
                row.append(Paragraph(cell_text, styles['Normal']))
                if len(row) == 2:
                    sig_data.append(row)
                    row = []
            if row:
                sig_data.append(row)

            if sig_data:
                table = Table(sig_data, colWidths=[250, 250])
                table.setStyle(TableStyle([
                    ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                    ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ]))
                story.append(table)

            doc.build(story)
            logger.info(f"PDF saved successfully: {file_path}")
            return file_path

        except Exception as e:
            logger.error(f"PDF creation failed: {str(e)}")
            raise DocumentCreationError(f"PDF generation failed: {str(e)}")